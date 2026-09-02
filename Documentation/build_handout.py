from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Documentation" / "ProcessDataAI-POC-Developer-Handout.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CODE_BG = "F6F8FA"
CAUTION = "FFF4CE"
SUCCESS = "EAF4EA"
WHITE = "FFFFFF"
MUTED = "5B6573"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("20262E")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 28, INK, 0, 8),
        ("Subtitle", 13, MUTED, 0, 16),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading")

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string("24292F")
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.0

    file_ref = styles.add_style("File Reference", WD_STYLE_TYPE.PARAGRAPH)
    file_ref.font.name = "Consolas"
    file_ref.font.size = Pt(9)
    file_ref.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    file_ref.paragraph_format.space_after = Pt(4)

    note = styles.add_style("Callout Text", WD_STYLE_TYPE.PARAGRAPH)
    note.font.name = "Calibri"
    note.font.size = Pt(10.5)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.line_spacing = 1.15


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.style = "Callout Text"
    run = p.add_run(title + "  ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code(doc, code, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(caption)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_BG)
    cell.text = ""
    for idx, line in enumerate(code.strip("\n").splitlines()):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.style = "Code Block"
        p.add_run(line if line else " ")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_file_link(doc, label, rel_path):
    p = doc.add_paragraph(style="File Reference")
    p.add_run(label + ": ").bold = True
    # Keep generated documents portable and avoid embedding the maintainer's
    # local filesystem path in the displayed text or hyperlink target.
    repository_path = rel_path.replace("\\", "/")
    add_hyperlink(p, repository_path, repository_path, DARK_BLUE)
    return p


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(INK)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            p.add_run(str(value))
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_source(doc, label, url):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_hyperlink(p, label, url)


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "ProcessDataAI  |  Developer Handout"
    header.style = doc.styles["Header"]
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    add_page_number(section.footer.paragraphs[0])

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("PROCESSDATAAI")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    title = doc.add_paragraph(style="Title")
    title.add_run("Document Ingestion and\nSemantic Search POC")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("A developer handout for understanding, running, rebuilding, and extending the application")
    add_callout(
        doc,
        "At a glance",
        "A .NET 10 console application reads local PDFs, normalizes extracted text into Microsoft Data Ingestion documents, chunks that content, generates embeddings through Azure OpenAI or Ollama, stores vectors in memory, and returns the three most semantically similar chunks for a question.",
        LIGHT_BLUE,
    )
    add_table(doc, ["Audience", "Purpose"], [
        ("Developer new to the Microsoft AI ingestion stack", "Understand the design and reproduce the implementation without prior library knowledge."),
        ("POC maintainer", "Operate the application, interpret logs, diagnose ingestion failures, and identify production-hardening work."),
    ], [2600, 6760])
    add_body(doc, "Project root: the local ProcessDataAI repository")
    add_body(doc, "Document version: July 2026")

    page_break(doc)
    add_heading(doc, "How to use this handout", 1)
    add_body(doc, "Read Sections 1-4 for the conceptual model and architecture. Sections 5-13 explain the implementation file by file. Sections 14-17 cover operation, troubleshooting, limitations, and a practical rebuild checklist.")
    add_heading(doc, "Contents", 2)
    add_numbers(doc, [
        "Application purpose and outcomes",
        "Core concepts",
        "Architecture and end-to-end flow",
        "Project structure",
        "Dependencies and why they exist",
        "Startup and dependency injection",
        "Environment configuration and validation",
        "Embedding provider abstraction",
        "PDF extraction and document normalization",
        "Chunking strategy",
        "Vector storage and chunk writing",
        "Ingestion orchestration",
        "Semantic search",
        "Running and validating the application",
        "Errors, logging, and troubleshooting",
        "Limitations and production evolution",
        "Reimplementation checklist and references",
    ])
    add_callout(doc, "Terminology", "In this handout, document means one PDF; section usually means one extracted PDF page; chunk means a token-bounded segment written to the vector store; record means the vector-store representation of one chunk.", LIGHT_GRAY)

    page_break(doc)
    add_heading(doc, "1. Application purpose and outcomes", 1)
    add_body(doc, "The POC demonstrates the smallest complete retrieval pipeline needed for semantic search over local documents. It is intentionally a console application and intentionally uses an in-memory vector store. This keeps the focus on ingestion abstractions, embeddings, and retrieval rather than persistence or web UI concerns.")
    add_heading(doc, "What the application does", 2)
    add_numbers(doc, [
        "Loads provider settings from a local .env file.",
        "Discovers every top-level PDF in Data/.",
        "Extracts selectable text from each PDF with PdfPig; no OCR is performed.",
        "Converts each PDF into an IngestionDocument with page-based sections.",
        "Splits documents with Microsoft.Extensions.DataIngestion.DocumentTokenChunker.",
        "Generates embeddings through Microsoft.Extensions.AI.",
        "Writes chunk content and vectors to Microsoft Semantic Kernel's InMemoryVectorStore connector.",
        "Accepts an interactive or one-shot question and returns the top three matches.",
    ])
    add_heading(doc, "What the application deliberately does not do", 2)
    add_bullets(doc, [
        "No OCR, so image-only or scanned PDFs are skipped with a clear error.",
        "No persistent database; all vectors disappear when the process exits.",
        "No answer generation. The output is retrieval evidence: score, document name, and chunk content.",
        "No web service, authentication, authorization, document upload, or multi-user isolation.",
    ])

    add_heading(doc, "2. Core concepts", 1)
    add_table(doc, ["Concept", "Meaning in this POC"], [
        ("Ingestion document", "A format-neutral Microsoft object containing sections and elements. It separates PDF parsing from downstream chunking."),
        ("Chunk", "A bounded piece of text suitable for embedding and retrieval. Chunks overlap so context near a boundary is not lost."),
        ("Embedding", "A numeric vector representing semantic meaning. Similar text tends to occupy nearby positions in vector space."),
        ("Vector store", "A collection that stores content plus vectors and supports nearest-neighbor search."),
        ("Cosine similarity", "The configured comparison method. Higher result scores indicate closer semantic alignment."),
        ("Provider abstraction", "The application consumes IEmbeddingGenerator, while startup chooses Azure OpenAI or Ollama."),
    ], [2100, 7260])
    add_callout(doc, "Key design idea", "The reader knows about PDFs, the chunker knows about ingestion documents, and the vector writer knows about chunks. Each stage depends on an abstraction rather than on the internal details of the previous technology.", SUCCESS)

    page_break(doc)
    add_heading(doc, "3. Architecture and end-to-end flow", 1)
    add_body(doc, "The application is a linear pipeline with two uses of the embedding generator: first during ingestion to vectorize chunks, then during search to vectorize the user's query.")
    add_table(doc, ["Stage", "Input", "Component", "Output"], [
        ("1. Configure", ".env", "EnvFile + AiOptionsValidator", "Validated AiOptions"),
        ("2. Select provider", "AiOptions", "EmbeddingGeneratorFactory", "IEmbeddingGenerator<string, Embedding<float>>"),
        ("3. Discover", "Data/", "DocumentSearchService", "Ordered FileInfo[]"),
        ("4. Read", "PDF stream", "PdfPigDocumentReader", "IngestionDocument"),
        ("5. Chunk", "IngestionDocument", "DocumentTokenChunker", "IAsyncEnumerable<IngestionChunk<string>>"),
        ("6. Embed + write", "Chunks", "VectorStoreWriter", "In-memory vector records"),
        ("7. Search", "Question text", "VectorStoreCollection.SearchAsync", "Top 3 scored records"),
    ], [1200, 1550, 3150, 3460])
    add_heading(doc, "Control flow", 2)
    add_code(doc, """Program.cs
  -> load .env
  -> build DI container
  -> resolve DocumentSearchService
  -> IngestAsync(Data/)
       -> probe embedding dimension
       -> read PDF -> normalize -> chunk -> embed -> store
  -> SearchAsync(question)
       -> embed question -> vector search -> print top 3
  -> dispose pipeline and vector store""", "Figure 1. Application lifecycle")
    add_heading(doc, "Why dependency injection matters", 2)
    add_body(doc, "Program.cs composes the application but does not implement PDF extraction, provider selection, ingestion, or search. This makes each responsibility testable and replaceable. For example, a future native Microsoft PDF reader could replace PdfPigDocumentReader without changing the chunker or search service.")

    add_heading(doc, "4. Project structure", 1)
    add_table(doc, ["Area", "File", "Responsibility"], [
        ("Entry point", "Program.cs", "Host, logging, DI, application mode, lifecycle."),
        ("Configuration", "Configuration/EnvFile.cs", "Minimal .env parsing with helpful format/missing-file errors."),
        ("Configuration", "Configuration/AiOptions.cs", "Strongly typed provider settings."),
        ("Configuration", "Configuration/AiOptionsValidator.cs", "Provider-specific required fields and endpoint validation."),
        ("Services", "Services/EmbeddingGeneratorFactory.cs", "Azure or Ollama embedding client creation."),
        ("Ingestion", "Ingestion/PdfPigDocumentReader.cs", "PDF extraction and normalization."),
        ("Ingestion", "Ingestion/CountingVectorStoreWriter.cs", "VectorStoreWriter adapter that counts chunks."),
        ("Services", "Services/DocumentSearchService.cs", "Document discovery, pipeline composition, ingestion, and search."),
        ("Project", "ProcessDataAI.csproj", "Target framework and NuGet dependencies."),
    ], [1350, 3050, 4960])

    page_break(doc)
    add_heading(doc, "5. Dependencies and why they exist", 1)
    add_table(doc, ["Package", "Version", "Role"], [
        ("Microsoft.Extensions.DataIngestion", "10.7.0-preview.1.26309.5", "Document model, chunker, pipeline, result, and vector writer."),
        ("Microsoft.Extensions.AI.OpenAI", "10.7.0", "Adapts OpenAI SDK embedding clients to IEmbeddingGenerator."),
        ("Microsoft.Extensions.VectorData", "Transitive", "Common vector-store collection and search abstractions."),
        ("Microsoft.SemanticKernel.Connectors.InMemory", "1.74.0-preview", "Concrete in-memory VectorStore implementation."),
        ("Azure.AI.OpenAI", "2.1.0", "AzureOpenAIClient for Azure deployments."),
        ("PdfPig", "0.1.12", "Local PDF text extraction without OCR."),
        ("Microsoft.ML.Tokenizers.Data.O200kBase", "2.0.0", "Tokenizer data used by the token chunker."),
        ("Microsoft.Extensions.Hosting", "10.0.9", "Host, DI, application lifetime, and configuration composition."),
        ("Microsoft.Extensions.Logging.Console", "10.0.9", "Structured console logging."),
        ("Microsoft.Bcl.Memory", "10.0.9", "Explicit patched transitive version to avoid a vulnerability warning."),
    ], [3750, 2450, 3160])
    add_file_link(doc, "Package references", "ProcessDataAI.csproj")
    add_callout(doc, "Preview packages", "Microsoft.Extensions.DataIngestion and the in-memory connector are preview dependencies. A production implementation should pin versions, test upgrades deliberately, and expect API changes.", CAUTION)

    add_heading(doc, "6. Startup and dependency injection", 1)
    add_body(doc, "Program.cs owns composition. It loads .env before building the host, adds the values as an in-memory configuration source, configures single-line timestamped logging, registers validated options and singleton services, starts the host, ingests Data/, and then chooses one-shot or interactive search.")
    add_code(doc, """HostApplicationBuilder builder = Host.CreateApplicationBuilder(args);
builder.Configuration.AddInMemoryCollection(envValues);

builder.Services.AddSingleton<PdfPigDocumentReader>();
builder.Services.AddSingleton<EmbeddingGeneratorFactory>();
builder.Services.AddSingleton<IEmbeddingGenerator<string, Embedding<float>>>(
    services => services.GetRequiredService<EmbeddingGeneratorFactory>().Create());
builder.Services.AddSingleton<DocumentSearchService>();""", "Core DI registration")
    add_body(doc, "Singleton lifetime is appropriate for this POC because one process owns one provider client, one in-memory store, and one ingestion/search session. The host disposes registered disposable services during shutdown.")
    add_file_link(doc, "Startup implementation", "Program.cs")

    page_break(doc)
    add_heading(doc, "7. Environment configuration and validation", 1)
    add_body(doc, "Provider selection is data-driven. Switching between Azure and Ollama requires only an .env change; the downstream pipeline always receives the same IEmbeddingGenerator abstraction.")
    add_code(doc, """AI_PROVIDER=Azure

AZURE_OPENAI_EMBEDDING_ENDPOINT=
AZURE_OPENAI_CHAT_ENDPOINT=
AZURE_OPENAI_API_KEY=
AZURE_OPENAI_EMBEDDING_MODEL=

OLLAMA_EMBEDDING_ENDPOINT=http://localhost:11434/v1
OLLAMA_CHAT_ENDPOINT=http://localhost:11434/v1
OLLAMA_EMBEDDING_MODEL=nomic-embed-text""", ".env.example")
    add_heading(doc, "Configuration lifecycle", 2)
    add_numbers(doc, [
        "EnvFile.Load checks that .env exists and parses non-empty, non-comment KEY=VALUE lines.",
        "Program.cs maps configuration keys into AiOptions and nested Azure/Ollama option objects.",
        "AiOptionsValidator checks AI_PROVIDER and validates only the selected provider's required fields.",
        "ValidateOnStart causes invalid configuration to fail before ingestion begins.",
    ])
    add_callout(doc, "Secret handling", "The real .env is ignored by Git. .env.example documents keys but must not contain credentials. For a deployed service, prefer a managed secret store or environment injection rather than a local file.", CAUTION)
    add_file_link(doc, "Example configuration", ".env.example")
    add_file_link(doc, "Parser", "Configuration/EnvFile.cs")
    add_file_link(doc, "Typed settings", "Configuration/AiOptions.cs")
    add_file_link(doc, "Validation", "Configuration/AiOptionsValidator.cs")

    add_heading(doc, "8. Embedding provider abstraction", 1)
    add_body(doc, "EmbeddingGeneratorFactory is the only class that knows how to create provider-specific clients. Both branches return IEmbeddingGenerator<string, Embedding<float>>, which is the Microsoft.Extensions.AI contract consumed everywhere else.")
    add_code(doc, """if (settings.Provider.Equals(AiOptions.AzureProvider,
    StringComparison.OrdinalIgnoreCase))
{
    var client = new AzureOpenAIClient(
        new Uri(azure.EmbeddingEndpoint),
        new AzureKeyCredential(azure.ApiKey));
    return client.GetEmbeddingClient(azure.EmbeddingModel)
        .AsIEmbeddingGenerator();
}

var options = new OpenAIClientOptions {
    Endpoint = new Uri(ollama.EmbeddingEndpoint)
};
var client = new OpenAIClient(new ApiKeyCredential("not-required"), options);
return client.GetEmbeddingClient(ollama.EmbeddingModel)
    .AsIEmbeddingGenerator();""", "Provider selection, shortened")
    add_body(doc, "Ollama exposes an OpenAI-compatible /v1 endpoint. The OpenAI SDK requires a credential object, so the factory supplies a placeholder string; Ollama ignores it. No real Ollama API key is read or required.")
    add_file_link(doc, "Provider factory", "Services/EmbeddingGeneratorFactory.cs")

    page_break(doc)
    add_heading(doc, "9. PDF extraction and document normalization", 1)
    add_body(doc, "Microsoft.Extensions.DataIngestion defines IngestionDocumentReader but does not parse these local PDFs by itself. PdfPigDocumentReader supplies the format-specific adapter. It opens the incoming stream, extracts each page in content order, and translates non-empty page text into Microsoft's neutral ingestion document model.")
    add_code(doc, """using PdfDocument pdf = PdfDocument.Open(source);
var document = new IngestionDocument(identifier);

foreach (var page in pdf.GetPages())
{
    string text = ContentOrderTextExtractor.GetText(page).Trim();
    if (string.IsNullOrWhiteSpace(text)) continue;

    var section = new IngestionDocumentSection();
    section.Elements.Add(
        new IngestionDocumentParagraph(text) { PageNumber = page.Number });
    document.Sections.Add(section);
}""", "Core reader logic")
    add_heading(doc, "Why ContentOrderTextExtractor is used", 2)
    add_body(doc, "PDFs store drawing instructions, not necessarily reading-order paragraphs. PdfPig's content-order extractor is a better default than directly reading page.Text because it attempts to reconstruct human reading order.")
    add_heading(doc, "Image-only PDFs", 2)
    add_body(doc, "If every page produces empty text, the reader throws InvalidDataException. The pipeline captures that failure as an IngestionResult, logs it, skips the document, and continues with other PDFs. This is why the two Jam-export PDFs produce error logs but do not stop a successful EmployeeHandbook search.")
    add_callout(doc, "No OCR by design", "A PDF can display words while containing only an image. Such a document has no selectable text for PdfPig to extract. OCR is a separate capability and is intentionally outside this POC.", CAUTION)
    add_file_link(doc, "PDF reader", "Ingestion/PdfPigDocumentReader.cs")

    add_heading(doc, "10. Chunking strategy", 1)
    add_body(doc, "The POC uses DocumentTokenChunker with the o200k tokenizer. HeaderChunker was not selected because raw PDF extraction does not reliably identify semantic heading levels. Inventing headings from layout would add fragile heuristics to a minimal sample.")
    add_code(doc, """var chunkerOptions = new IngestionChunkerOptions(
    TiktokenTokenizer.CreateForModel("gpt-4o"))
{
    MaxTokensPerChunk = 500,
    OverlapTokens = 75
};

IngestionChunker<string> chunker =
    new DocumentTokenChunker(chunkerOptions);""", "Chunker configuration")
    add_table(doc, ["Setting", "Value", "Reason"], [
        ("Maximum", "500 tokens", "Small enough for focused retrieval; large enough for useful paragraph context."),
        ("Overlap", "75 tokens", "Repeats boundary context so an answer near a split is less likely to be lost."),
        ("Tokenizer", "o200k / gpt-4o mapping", "Provides deterministic token counting independent of the embedding provider."),
    ], [2100, 1700, 5560])

    page_break(doc)
    add_heading(doc, "11. Vector storage and chunk writing", 1)
    add_body(doc, "The Microsoft VectorStoreWriter requires the embedding dimension when it creates its collection schema. Dimensions differ by model, so the service generates a small probe embedding and reads probe.Length rather than hardcoding a model-specific number.")
    add_code(doc, """ReadOnlyMemory<float> probe = await embeddingGenerator.GenerateVectorAsync(
    "embedding dimension probe",
    cancellationToken: cancellationToken);

_vectorStore = new InMemoryVectorStore(new()
{
    EmbeddingGenerator = embeddingGenerator
});

_writer = new CountingVectorStoreWriter(
    _vectorStore,
    probe.Length,
    new VectorStoreWriterOptions
    {
        CollectionName = "documents",
        DistanceFunction = DistanceFunction.CosineSimilarity
    });""", "Dimension probe and store initialization")
    add_heading(doc, "Why the store receives the embedding generator", 2)
    add_body(doc, "The vector-data layer can automatically generate vectors from text when writing and when searching with a text query. That lets the application call SearchAsync(query, top: 3) instead of manually generating a query vector.")
    add_heading(doc, "CountingVectorStoreWriter", 2)
    add_body(doc, "VectorStoreWriter does not expose a simple total-chunk counter. The adapter derives from IngestionChunkWriter<string>, streams every chunk to the real writer, increments ChunkCount, and exposes the underlying collection for search. It does not buffer all chunks in memory before forwarding them.")
    add_file_link(doc, "Counting adapter", "Ingestion/CountingVectorStoreWriter.cs")

    add_heading(doc, "12. Ingestion orchestration", 1)
    add_body(doc, "DocumentSearchService is the application service coordinating filesystem checks, provider verification, vector-store construction, chunker construction, pipeline execution, result handling, metrics, and lifecycle.")
    add_code(doc, """_pipeline = new IngestionPipeline<string>(
    reader,
    chunker,
    _writer,
    loggerFactory: loggerFactory);

await foreach (IngestionResult result in
    _pipeline.ProcessAsync(pdfs, cancellationToken))
{
    if (!result.Succeeded)
    {
        logger.LogError(result.Exception,
            "Failed to ingest PDF {DocumentName}", result.DocumentId);
        continue;
    }
    successfulDocuments++;
}""", "Pipeline composition and per-document handling")
    add_body(doc, "A failed PDF does not abort the complete batch. The service only fails the application when zero PDFs were successfully ingested. This is a deliberate partial-success policy for a directory ingestion job.")
    add_callout(doc, "Pipeline lifetime", "The pipeline is stored as a service field rather than a local using variable. Disposing it immediately after IngestAsync would dispose the writer before SearchAsync runs. DocumentSearchService disposes the pipeline and vector store at application shutdown.", SUCCESS)

    page_break(doc)
    add_heading(doc, "13. Semantic search", 1)
    add_body(doc, "After at least one document succeeds, SearchAsync uses the writer's dynamic vector-store collection. The collection embeds the query, compares it with stored chunk embeddings using cosine similarity, and yields up to three VectorSearchResult records.")
    add_code(doc, """await foreach (var result in
    _writer.VectorStoreCollection.SearchAsync(
        query, top: 3, cancellationToken: cancellationToken))
{
    string documentId = GetString(result.Record,
        "documentid", "document_id", "documentId", "DocumentId")
        ?? "unknown";
    string content = GetString(result.Record, "content", "Content")
        ?? string.Empty;

    Console.WriteLine($"Score: {result.Score:F4}");
    Console.WriteLine($"Document: {Path.GetFileName(documentId)}");
    Console.WriteLine($"Content: {content}");
}""", "Top-three search and output")
    add_heading(doc, "Interpreting the score", 2)
    add_body(doc, "The score ranks results within the same search. It should not be treated as a universal correctness probability. Thresholds, if introduced later, should be calibrated on representative questions and documents.")
    add_heading(doc, "Why the dynamic record uses multiple key spellings", 2)
    add_body(doc, "VectorStoreWriter exposes a dynamic Dictionary<string, object?> schema. The current preview produces lower-case documentid and content keys. GetString accepts several variants to tolerate minor naming changes across preview versions.")
    add_file_link(doc, "Ingestion and search service", "Services/DocumentSearchService.cs")

    add_heading(doc, "14. Running and validating the application", 1)
    add_heading(doc, "Prerequisites", 2)
    add_bullets(doc, [
        ".NET 10 SDK.",
        "At least one text-bearing PDF in Data/.",
        "An Ollama embedding model available at the configured endpoint, or a deployed Azure OpenAI embedding model.",
        "A .env file copied from .env.example and completed for the selected provider.",
    ])
    add_code(doc, """dotnet restore
dotnet build
dotnet run

# One-shot query
dotnet run -- --query "How many vacation days per calendar year?""" , "Build and run")
    add_heading(doc, "Expected one-shot output", 2)
    add_code(doc, """Score: 0.7703
Document: EmployeeHandbook.pdf
Content: Employee Handbook
Vacation Policy
Full-time employees receive 25 paid vacation days per calendar year.
Unused vacation days may be carried over up to a maximum of five days.""", "Example result; score can vary by model")
    add_body(doc, "Without --query, the application repeatedly prompts Question (or 'exit'): and reuses the same in-memory collection until the user exits.")

    page_break(doc)
    add_heading(doc, "15. Errors, logging, and troubleshooting", 1)
    add_table(doc, ["Symptom", "Cause", "Action"], [
        (".env not found", "Startup configuration file is missing.", "Copy .env.example to .env and configure one provider."),
        ("AI_PROVIDER must be Azure or Ollama", "Unsupported or misspelled provider value.", "Use exactly Azure or Ollama; comparison is case-insensitive."),
        ("Missing required provider configuration", "A selected-provider field is empty.", "Complete endpoint, model/deployment, and Azure key when applicable."),
        ("Data directory was not found", "Application started from the wrong working directory or Data/ is missing.", "Run from ProcessDataAI or create Data/."),
        ("Data contains no PDF files", "No top-level *.pdf files were discovered.", "Add text-bearing PDFs; subdirectories are not searched."),
        ("Contains no extractable text", "The PDF is image-only, scanned, or blank.", "Use a text PDF, remove it, or add an OCR stage in a future version."),
        ("Embedding generation failed", "Endpoint, model, credentials, network, or Ollama availability problem.", "Check .env, confirm the model exists, and test endpoint reachability."),
        ("No PDFs could be ingested", "Every discovered PDF failed reading or writing.", "Review the preceding per-document errors."),
        ("Semantic search failed", "Query embedding or vector search failed.", "Check provider health and ensure ingestion completed."),
    ], [2900, 3160, 3300])
    add_heading(doc, "Why PDF errors appear twice", 2)
    add_body(doc, "The Microsoft IngestionPipeline logs the exception when reading fails. DocumentSearchService then logs the unsuccessful IngestionResult. Both refer to the same failure. They are fail-level log entries, but they are non-fatal when another document succeeds.")
    add_heading(doc, "Useful success logs", 2)
    add_bullets(doc, [
        "Loaded AI provider confirms provider selection.",
        "Discovered N PDF file(s) confirms directory scanning.",
        "Extracted text from N page(s) confirms PdfPig normalization.",
        "Ingestion completed reports successful documents, chunk count, and elapsed time.",
        "Executing semantic search confirms query execution.",
    ])

    add_heading(doc, "16. Limitations and production evolution", 1)
    add_table(doc, ["POC limitation", "Production direction"], [
        ("In-memory vectors", "Use a supported persistent vector database and define migration/backup behavior."),
        ("Reingests on every run", "Add document fingerprints, incremental ingestion, and deletion/update handling."),
        ("No OCR", "Insert an OCR/document-intelligence reader for scanned documents, subject to cost and privacy requirements."),
        ("One process, one user", "Expose a service API with authentication, tenant isolation, quotas, and cancellation."),
        ("No generated answer", "Optionally add a chat model that answers only from retrieved chunks and includes citations."),
        ("No retrieval evaluation", "Build a test set of questions and expected documents; measure recall, ranking, and latency."),
        ("Dynamic record schema", "Use an explicit typed vector record when the preview API and target provider allow it."),
        ("Preview packages", "Pin versions, monitor release notes, add integration tests, and budget for API migrations."),
        ("Local .env secrets", "Use managed identity, workload identity, Key Vault, or deployment-managed secrets."),
    ], [3150, 6210])

    page_break(doc)
    add_heading(doc, "17. Reimplementation checklist and references", 1)
    add_heading(doc, "Rebuild the POC from an existing console project", 2)
    add_numbers(doc, [
        "Add the required NuGet packages and remove any persistent database connector not needed by the POC.",
        "Create .env.example, ignore .env, load it before host construction, and validate typed options at startup.",
        "Return one IEmbeddingGenerator implementation from a provider factory selected by AI_PROVIDER.",
        "Implement IngestionDocumentReader for the source format; create one IngestionDocument per PDF and page-based sections.",
        "Choose a Microsoft Data Ingestion chunker. Use token chunking when source headings are not trustworthy.",
        "Generate a probe embedding to discover dimensions, then construct InMemoryVectorStore and VectorStoreWriter.",
        "Compose IngestionPipeline and process all PDFs while allowing partial success.",
        "Keep pipeline/writer/store alive through the search phase and dispose them at shutdown.",
        "Call SearchAsync with text and top: 3; print score, source name, and content.",
        "Verify restore, warning-free build, PDF discovery, extraction, chunk count, embeddings, and a known-answer search.",
    ])
    add_heading(doc, "Project file links", 2)
    for label, path in (
        ("Entry point", "Program.cs"),
        ("Project dependencies", "ProcessDataAI.csproj"),
        ("Environment example", ".env.example"),
        ("Typed configuration", "Configuration/AiOptions.cs"),
        ("Configuration validator", "Configuration/AiOptionsValidator.cs"),
        ("Environment parser", "Configuration/EnvFile.cs"),
        ("Provider factory", "Services/EmbeddingGeneratorFactory.cs"),
        ("PDF reader", "Ingestion/PdfPigDocumentReader.cs"),
        ("Counting writer", "Ingestion/CountingVectorStoreWriter.cs"),
        ("Ingestion/search service", "Services/DocumentSearchService.cs"),
        ("Quick README", "README.md"),
    ):
        add_file_link(doc, label, path)

    add_heading(doc, "Official and package references", 2)
    add_source(doc, "Microsoft Data Ingestion quickstart", "https://learn.microsoft.com/en-us/dotnet/ai/quickstarts/process-data")
    add_source(doc, "Microsoft Vector Stores guide", "https://learn.microsoft.com/en-us/dotnet/ai/vector-stores/how-to/use-vector-stores")
    add_source(doc, "Microsoft.Extensions.DataIngestion API", "https://learn.microsoft.com/en-us/dotnet/api/microsoft.extensions.dataingestion")
    add_source(doc, "Microsoft.Extensions.DataIngestion NuGet", "https://www.nuget.org/packages/Microsoft.Extensions.DataIngestion/")
    add_source(doc, "In-memory vector connector NuGet", "https://www.nuget.org/packages/Microsoft.SemanticKernel.Connectors.InMemory")
    add_source(doc, "PdfPig NuGet and usage", "https://www.nuget.org/packages/PdfPig")

    add_heading(doc, "Final mental model", 2)
    add_callout(doc, "Remember the boundary", "PdfPig turns a PDF into Microsoft's neutral document model. Microsoft Data Ingestion turns that document into chunks. Microsoft.Extensions.AI turns text into vectors. Microsoft.Extensions.VectorData and InMemoryVectorStore store and retrieve those vectors. DocumentSearchService wires the stages together.", SUCCESS)

    doc.core_properties.title = "ProcessDataAI Document Ingestion and Semantic Search POC - Developer Handout"
    doc.core_properties.subject = "Implementation overview and developer guide"
    doc.core_properties.author = "ProcessDataAI Project"
    doc.core_properties.keywords = "Data Ingestion, Vector Search, Microsoft.Extensions.AI, PdfPig, Ollama, Azure OpenAI"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
